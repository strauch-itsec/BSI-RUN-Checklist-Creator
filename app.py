from flask import Flask, render_template, request, send_file, jsonify, after_this_request, session
import json
import os
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt

app = Flask(__name__)

DOWNLOADS_DIR = "downloads"
os.makedirs(DOWNLOADS_DIR, exist_ok=True)   # Ensure the downloads directory exists

# Path to the secret key file
SECRET_KEY_FILE = ".secret_key"

def generate_secret_key():
    """Generate a new secret key."""
    return os.urandom(24)

def load_or_create_secret_key():
    """Load the secret key from a file or create a new one if it doesn't exist."""
    if os.path.exists(SECRET_KEY_FILE):
        # Load the existing secret key
        with open(SECRET_KEY_FILE, "rb") as f:
            return f.read()
    else:
        # Generate a new secret key
        secret_key = generate_secret_key()
        with open(SECRET_KEY_FILE, "wb") as f:
            f.write(secret_key)
        return secret_key

# Set the secret key for the Flask app
app.secret_key = load_or_create_secret_key()

# JSON-Datei laden mit Fehlerbehandlung
def load_json():
    try:
        if not os.path.exists("run_Anhang.json"):
            return {}

        with open("run_Anhang.json", "r", encoding="utf-8") as file:
            return json.load(file)
    except (json.JSONDecodeError, IOError) as e:
        print(f"Fehler beim Laden der JSON-Datei: {e}")
        return {}

def get_filtered_data():
    """
    Retrieves and filters data based on the selected category and grades stored in the session.
    :return: A tuple (filtered_data, error_message). If no error, error_message will be None.
    """
    # Load data from the JSON file
    data = load_json()

    # Retrieve selected category and grades from the session
    selected_category = session.get("selected_category", "Unknown")
    selected_grades = session.get("selected_grades", [])
    
    # Validate session data
    if selected_category == "Unknown" or selected_category not in data:
        return None, "Ungültige Kategorie"
    if not selected_grades:
        return None, "Keine Umsetzungsgrade ausgewählt"

    # Filter the data based on the selected category and grades
    filtered_data = [req for req in data[selected_category] if req["Umsetzungsgrad"] in selected_grades]
    return filtered_data, None        

# Startseite mit Filtermöglichkeiten
@app.route('/')
def index():
    data = load_json()
    categories = list(data.keys()) if data else []  # Themenbereiche
    max_items = sum(len(items) for items in data.values()) if data else 0
    script_name = request.environ.get('SCRIPT_NAME', '')  # Get the SCRIPT_NAME
    return render_template("index.html", categories=categories, max_items=max_items, script_name=script_name)
    
# Gefilterte Anforderungen anzeigen
@app.route('/filter', methods=['POST'])
def filter_data():
    data = load_json()

    # Define a whitelist for valid grades
    valid_grades = {"UG2", "UG3", "UG4", "UG5"}
    
    selected_grades = request.form.getlist("umsetzungsgrad")
    selected_category = request.form.get("category", "").strip()

    # Validate the category
    if not selected_category or selected_category not in data:
        return jsonify({"error": "Ungültige Kategorie"}), 400

    # Validate the grades against the whitelist
    if not all(grade in valid_grades for grade in selected_grades):
        return jsonify({"error": "Ungültige Umsetzungsgrade"}), 400

    # Store selected values in the session
    session['selected_category'] = selected_category
    session['selected_grades'] = selected_grades    

    filtered = [req for req in data[selected_category] if req["Umsetzungsgrad"] in selected_grades]
    return jsonify(filtered)

def create_word_document(filtered_data, full_description=False):
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1  # Querformat
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    doc.add_heading("Checkliste", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Spaltenbreite einstellen
    table.columns[0].width = Inches(1)  # Schmale Spalte
    table.columns[2].width = Inches(1)  # Schmale Spalte

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Anforderungsnr."
    hdr_cells[1].text = "Name"
    hdr_cells[2].text = "Umsetzungsgrad"
    hdr_cells[3].text = "Status"

    last_integer = None
    for item in filtered_data:
        current_number = item['Anforderungsnummer']
        
        # Prüfen, ob die Anforderungsnummer einen Punkt enthält
        if '.' in str(current_number):
            current_integer = int(str(current_number).split('.')[0])
            if current_integer != last_integer:
                # Neue Zeile für die Integer-Nummer und den Namen hinzufügen
                integer_row = table.add_row().cells
                integer_row[0].text = str(current_integer)
                paragraph = integer_row[1].paragraphs[0]
                run = paragraph.add_run(item['Name'])
                run.bold = True
                integer_row[2].text = ""
                integer_row[3].text = ""
                last_integer = current_integer

            # Zeile für den Unterpunkt hinzufügen
            row_cells = table.add_row().cells
            row_cells[0].text = str(current_number)
            paragraph = row_cells[1].paragraphs[0]
            if full_description:
                paragraph.add_run(item['Beschreibung'])
            else:
                first_sentence = item['Beschreibung'].split('.')[0] + '.'
                paragraph.add_run(first_sentence)
            row_cells[2].text = item['Umsetzungsgrad']
            row_cells[3].text = ""
        else:
            # Zeile für die Integer-Nummer ohne Unterpunkte
            row_cells = table.add_row().cells
            row_cells[0].text = str(current_number)
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run(item['Name'])
            run.bold = True
            if full_description:
                paragraph.add_run(f"\n{item['Beschreibung']}")
            else:
                first_sentence = item['Beschreibung'].split('.')[0] + '.'
                paragraph.add_run(f"\n{first_sentence}")
            row_cells[2].text = item['Umsetzungsgrad']
            row_cells[3].text = ""
    # Add footer
    section = doc.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "Erstellt mit https://github.com/strauch-itsec/BSI-RUN-Checklist-Creator                   Disclaimer: Es wird keine Gewähr für die Richtigkeit der Daten übernommen."
    footer_paragraph.style.font.size = Pt(8)
    
    return doc

class CustomPDF(FPDF):
    def footer(self):
        """Define a footer that appears on every page."""
        self.set_y(-15)  # Position the footer 15 mm from the bottom
        self.set_font("DejaVu", size=8)
        self.set_text_color(0, 0, 255)
        self.cell(0, 10, "BSI-RUN-Checklist-Creator", 0, 0, "C", link="https://github.com/strauch-itsec/BSI-RUN-Checklist-Creator")
        self.set_text_color(0, 0, 0)
        self.ln(5)
        self.cell(0, 10, "Disclaimer: Es wird keine Gewähr für die Richtigkeit der Daten übernommen.", 0, 0, "C")

@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    try:
        # Get filtered data
        filtered_data, error = get_filtered_data()
        if error:
            return jsonify({"error": error}), 400

        # Retrieve session data for filename generation
        selected_category = session.get("selected_category", "Unknown").replace(" ", "-")
        selected_grades = session.get("selected_grades", [])
        grades_str = "-".join(selected_grades)
        filename = f"RUN-checklist-{selected_category}-{grades_str}.pdf"

        # Use CustomPDF for PDF generation
        pdf = CustomPDF(orientation="L", unit="mm", format="A4")  # Landscape format
        pdf.add_page()

        # Add a UTF-8 compatible font (e.g., DejaVu Sans)
        font_path = os.path.join("fonts", "DejaVuSans.ttf")  # Ensure the font file is in the "fonts" folder
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=12)

        pdf.cell(0, 10, "Checkliste", ln=True, align="C")
        pdf.ln(10)

        # Header row
        pdf.set_font("DejaVu", size=10)
        pdf.cell(30, 10, "Anforderungsnr.", 1)
        pdf.cell(120, 10, "Name", 1)
        pdf.cell(35, 10, "Umsetzungsgrad", 1)
        pdf.cell(50, 10, "Status", 1)
        pdf.ln()

        # Data rows
        for item in filtered_data:
            name_text = item['Name']
            description = item['Beschreibung']
            pdf.cell(30, 10, str(item['Anforderungsnummer']), 1)
            pdf.cell(120, 10, name_text[:50] + ("..." if len(name_text) > 50 else ""), 1)
            pdf.cell(35, 10, item['Umsetzungsgrad'], 1)
            pdf.cell(50, 10, "", 1)  # Empty status field
            pdf.ln()

        # Save the PDF
        pdf_path = os.path.join(DOWNLOADS_DIR, filename)
        pdf.output(pdf_path)

        @after_this_request
        def cleanup(response):
            try:
                os.remove(pdf_path)
            except Exception:
                print(f"Error deleting file {pdf_path}")
            return response

        return send_file(pdf_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": f"Fehler beim Exportieren: PDF "}), 500
    
# Word-Export Variante 1
@app.route('/export/word_variante1', methods=['POST'])
def export_word_variante1():
    try:
        # Get filtered data
        filtered_data, error = get_filtered_data()
        if error:
            return jsonify({"error": error}), 400

        # Retrieve session data for filename generation
        selected_category = session.get("selected_category", "Unknown").replace(" ", "-")
        selected_grades = session.get("selected_grades", [])
        grades_str = "-".join(selected_grades)
        filename = f"RUN-checklist-{selected_category}-{grades_str}-Variante1.docx"

        doc = create_word_document(filtered_data, full_description=False)
        doc_path = os.path.join(DOWNLOADS_DIR, filename)
        doc.save(doc_path)

        @after_this_request
        def cleanup(response):
            try:
                os.remove(doc_path)
            except Exception:
                print(f"Error deleting file {doc_path}")
            return response

        return send_file(doc_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": f"Fehler beim Exportieren: Word Variante 1"}), 500

# Word-Export Variante 2
@app.route('/export/word_variante2', methods=['POST'])
def export_word_variante2():
    try:
        # Get filtered data
        filtered_data, error = get_filtered_data()
        if error:
            return jsonify({"error": error}), 400

        # Retrieve session data for filename generation
        selected_category = session.get("selected_category", "Unknown").replace(" ", "-")
        selected_grades = session.get("selected_grades", [])
        grades_str = "-".join(selected_grades)
        filename = f"RUN-checklist-{selected_category}-{grades_str}-Variante2.docx"

        # Generate the filename dynamically
        grades_str = "-".join(selected_grades)
        filename = f"RUN-checklist-{selected_category}-{grades_str}-Variante2.docx"

        doc = create_word_document(filtered_data, full_description=True)
        doc_path = os.path.join(DOWNLOADS_DIR, filename)
        doc.save(doc_path)

        @after_this_request
        def cleanup(response):
            try:
                os.remove(doc_path)
            except Exception:
                print(f"Error deleting file {doc_path}")
            return response

        return send_file(doc_path, as_attachment=True)
    
    except Exception as e:
        return jsonify({"error": f"Fehler beim Exportieren: Word Variante 2"}), 500

if __name__ == '__main__':
    app.run(debug=False)
